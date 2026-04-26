import io
import os
import re
import tempfile
import traceback
from typing import List, Literal
from uuid import UUID

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import Response, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.application.dto.requests import GenerateCVRequest, ChatContextRequest, GeneratedCVUpdateRequest
from app.application.dto.responses import (
    ChatContextResponse,
    GeneratedCVListResponse,
    GeneratedCVResponse,
    GeneratedCVVersionResponse,
)
from app.application.use_cases.edit_generated_cv import EditGeneratedCVUseCase
from app.application.use_cases.generate_cv import GenerateCVUseCase
from app.application.use_cases.import_generated_cv import ImportGeneratedCVUseCase
from app.application.use_cases.chat_cv import ChatCVUseCase
from app.infrastructure.ai import ai_service_factory
from app.infrastructure.database.session import get_db_session
from app.infrastructure.database.repositories.generated_cv_repository import GeneratedCVRepository
from app.infrastructure.file_parsers.parsers import get_parser
from app.infrastructure.file_parsers.import_pipeline import (
    build_import_preview_payload,
    convert_pdf_to_docx,
)
from app.infrastructure.file_parsers.upload_validation import read_and_validate_upload
from app.presentation.auth_routes import get_current_user_id
from app.logger import get_logger

logger = get_logger("app.presentation.generated_cv_routes")

router = APIRouter(prefix="/generated-cvs", tags=["Generated CVs"])


def _strip_markdown_inline(text: str) -> str:
    content = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    content = content.replace("**", "").replace("*", "").replace("`", "")
    return content.strip()


def _markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            heading_level = min(len(heading_match.group(1)), 4)
            doc.add_heading(_strip_markdown_inline(heading_match.group(2)), level=heading_level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line) or re.match(r"^\d+\.\s+(.*)$", line)
        if bullet_match:
            doc.add_paragraph(_strip_markdown_inline(bullet_match.group(1)), style="List Bullet")
            continue

        doc.add_paragraph(_strip_markdown_inline(line))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _get_generated_content_payload(cv_entity) -> tuple[str, str]:
    content_data = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    output_format = content_data.get("format")

    if output_format not in {"markdown", "docx"}:
        if isinstance(content_data.get("markdown"), str):
            output_format = "markdown"
        else:
            output_format = "markdown"

    content = (
        content_data.get("content")
        or content_data.get("markdown")
        or ""
    )
    return output_format, content


def _build_export_filename(cv_entity, ext: str) -> str:
    job_title = (cv_entity.base_profile_data or {}).get("job_title") or "generated_cv"
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "_", str(job_title).strip().lower()).strip("_")
    if not normalized:
        normalized = "generated_cv"
    return f"{normalized[:60]}.{ext}"


async def _parse_uploaded_cv(file: UploadFile) -> dict[str, str]:
    settings = get_settings()
    file_bytes, upload_meta = await read_and_validate_upload(
        file,
        allowed_types={"pdf", "docx"},
        max_size_mb=settings.MAX_FILE_SIZE_MB,
        detail="Chỉ hỗ trợ file CV định dạng PDF hoặc DOCX",
    )
    normalized_name = upload_meta.filename.lower()
    ext = upload_meta.extension
    tmp_path = None
    converted_docx_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        if normalized_name.endswith(".pdf"):
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as converted:
                converted_docx_path = converted.name
            convert_pdf_to_docx(tmp_path, converted_docx_path)
            preview_payload = build_import_preview_payload(converted_docx_path)
        elif normalized_name.endswith(".docx"):
            preview_payload = build_import_preview_payload(tmp_path)
        else:
            parser = get_parser(upload_meta.filename)
            parsed_text = await parser.parse(tmp_path)
            preview_payload = {
                "markdown": parsed_text.strip(),
                "html": "",
            }
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Failed to import uploaded CV %s: %s\n%s", upload_meta.filename, str(exc), traceback.format_exc())
        raise HTTPException(status_code=400, detail=f"Không đọc được file CV: {exc}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        if converted_docx_path and os.path.exists(converted_docx_path):
            os.remove(converted_docx_path)

    if not preview_payload.get("markdown", "").strip():
        raise HTTPException(status_code=400, detail="Không trích xuất được nội dung từ file CV")

    return preview_payload


def _to_generated_cv_response(cv_entity) -> GeneratedCVResponse:
    return GeneratedCVResponse(
        id=cv_entity.id,
        conversation_id=cv_entity.conversation_id,
        version=cv_entity.version,
        parent_version_id=cv_entity.parent_version_id,
        status=cv_entity.status,
        target_jd_text=cv_entity.target_jd_text,
        base_profile_data=cv_entity.base_profile_data,
        generated_content=cv_entity.generated_content,
        created_at=cv_entity.created_at,
    )


def _to_generated_cv_version_response(cv_entity) -> GeneratedCVVersionResponse:
    return GeneratedCVVersionResponse(
        id=cv_entity.id,
        conversation_id=cv_entity.conversation_id,
        version=cv_entity.version,
        parent_version_id=cv_entity.parent_version_id,
        status=cv_entity.status,
        created_at=cv_entity.created_at,
    )


@router.post("/chat", response_model=ChatContextResponse)
async def chat_cv_generation(
    req: ChatContextRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Interact with CV AI chatbot."""
    cv_repo = GeneratedCVRepository(session)
    ai_service = ai_service_factory()

    try:
        messages = [{"role": msg.role, "content": msg.content} for msg in req.messages]
        current_cv = None
        if req.current_cv_id:
            current_cv = await cv_repo.get_by_id(req.current_cv_id)
            if not current_cv or current_cv.user_id != user_id:
                raise HTTPException(status_code=404, detail="Không tìm thấy phiên bản CV hiện tại")

        if current_cv:
            use_case = EditGeneratedCVUseCase(cv_repo, ai_service)
            reply, new_cv, _ = await use_case.execute(
                user_id=user_id,
                current_cv=current_cv,
                messages=messages,
                output_format=req.output_format,
            )
            cv_id = new_cv.id if new_cv else None
        else:
            use_case = ChatCVUseCase(cv_repo, ai_service)
            reply, cv_id = await use_case.execute(
                user_id=user_id,
                messages=messages,
                output_format=req.output_format,
                template_id=req.template_id,
            )

        if cv_id:
            await session.commit()
            
        return ChatContextResponse(
            reply=reply,
            generated_cv_id=cv_id
        )
    except HTTPException:
        await session.rollback()
        raise
    except Exception as e:
        await session.rollback()
        logger.error("Failed in chat interaction: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi giao tiếp với AI")

@router.post("/chat/stream")
async def chat_cv_generation_stream(
    req: ChatContextRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Interact with CV AI chatbot via Server-Sent Events (SSE)."""
    cv_repo = GeneratedCVRepository(session)
    ai_service = ai_service_factory()

    try:
        messages = [{"role": msg.role, "content": msg.content} for msg in req.messages]
        if req.current_cv_id:
            current_cv = await cv_repo.get_by_id(req.current_cv_id)
            if not current_cv or current_cv.user_id != user_id:
                raise HTTPException(status_code=404, detail="Không tìm thấy phiên bản CV hiện tại")

            use_case = EditGeneratedCVUseCase(cv_repo, ai_service)

            async def _edit_stream():
                import json

                try:
                    yield f"event: status\ndata: {json.dumps({'state': 'reasoning', 'label': 'AI đang phân tích yêu cầu chỉnh sửa...'})}\n\n"
                    reply, new_cv, next_content = await use_case.execute(
                        user_id=user_id,
                        current_cv=current_cv,
                        messages=messages,
                        output_format=req.output_format,
                    )
                    if reply:
                        yield f"event: chat_chunk\ndata: {json.dumps(reply)}\n\n"

                    if new_cv:
                        yield f"event: status\ndata: {json.dumps({'state': 'applying_edits', 'label': 'Đang áp thay đổi vào CV hiện tại...'})}\n\n"
                        await session.commit()
                        yield f"event: cv_chunk\ndata: {json.dumps(next_content)}\n\n"
                        yield f"event: status\ndata: {json.dumps({'state': 'saving_version', 'label': 'Đã lưu thành phiên bản CV mới.'})}\n\n"
                        yield f"event: cv_id\ndata: {json.dumps(str(new_cv.id))}\n\n"
                        yield f"event: status\ndata: {json.dumps({'state': 'done', 'label': 'Hoàn tất cập nhật CV.'})}\n\n"
                    else:
                        await session.rollback()
                        yield f"event: status\ndata: {json.dumps({'state': 'waiting_input', 'label': 'Mình cần thêm thông tin trước khi sửa CV.'})}\n\n"
                except Exception as exc:
                    await session.rollback()
                    logger.error("Failed to edit CV via chat stream: %s", str(exc), exc_info=True)
                    yield f"event: error\ndata: {json.dumps(str(exc))}\n\n"

            return StreamingResponse(
                _edit_stream(),
                media_type="text/event-stream",
            )

        use_case = ChatCVUseCase(cv_repo, ai_service)
        stream_generator = use_case.execute_stream(
            user_id=user_id,
            messages=messages,
            output_format=req.output_format,
            template_id=req.template_id,
        )
        return StreamingResponse(
            stream_generator,
            media_type="text/event-stream"
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to start chat streaming: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi bắt đầu stream AI")

@router.post("/", response_model=GeneratedCVResponse)
async def generate_cv(
    req: GenerateCVRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Generate a template CV via AI and save it."""
    
    cv_repo = GeneratedCVRepository(session)
    ai_service = ai_service_factory()
    use_case = GenerateCVUseCase(cv_repo, ai_service)
    
    try:
        cv_entity = await use_case.execute(
            user_id=user_id,
            job_title=req.job_title,
            jd_text=req.jd_text,
            level=req.level,
            output_format=req.output_format,
        )
        await session.commit()
        return _to_generated_cv_response(cv_entity)

    except HTTPException:
        await session.rollback()
        raise
    except Exception as e:
        await session.rollback()
        logger.error("Failed to generate CV: %s", str(e), exc_info=True)
        raise HTTPException(status_code=500, detail="Lỗi khi AI tạo CV mẫu")


@router.post("/import", response_model=GeneratedCVResponse, status_code=201)
async def import_generated_cv(
    cv_file: UploadFile = File(...),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Import an existing CV file and open it in the editable workspace."""
    cv_repo = GeneratedCVRepository(session)
    use_case = ImportGeneratedCVUseCase(cv_repo)

    logger.info(
        "Import generated CV request: user_id=%s, filename=%s, content_type=%s",
        user_id,
        cv_file.filename,
        cv_file.content_type,
    )

    try:
        preview_payload = await _parse_uploaded_cv(cv_file)
        cv_entity = await use_case.execute(
            user_id=user_id,
            filename=cv_file.filename or "uploaded_cv",
            parsed_content=preview_payload["markdown"],
            preview_html=preview_payload.get("html", ""),
        )
        await session.commit()
        logger.info("Imported CV saved successfully: cv_id=%s, filename=%s", cv_entity.id, cv_file.filename)
        return _to_generated_cv_response(cv_entity)
    except HTTPException:
        await session.rollback()
        raise
    except ValueError as exc:
        await session.rollback()
        raise HTTPException(status_code=400, detail=str(exc))
    except Exception as exc:
        await session.rollback()
        logger.error("Failed to import CV into workspace: %s", str(exc), exc_info=True)
        raise HTTPException(status_code=500, detail="Không thể import CV vào workspace")


@router.get("/", response_model=List[GeneratedCVListResponse])
async def list_generated_cvs(
    limit: int = 20,
    offset: int = 0,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """List user's generated CVs."""
    cv_repo = GeneratedCVRepository(session)
    cvs = await cv_repo.list_by_user_id(user_id, limit, offset)
    
    return [
        GeneratedCVListResponse(
            id=c.id,
            conversation_id=c.conversation_id,
            version=c.version,
            status=c.status,
            target_jd_text=c.target_jd_text,
            job_title=c.base_profile_data.get("job_title") if c.base_profile_data else None,
            level=c.base_profile_data.get("level") if c.base_profile_data else None,
            created_at=c.created_at,
        )
        for c in cvs
    ]


@router.get("/{cv_id}", response_model=GeneratedCVResponse)
async def get_generated_cv(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Get full generated CV."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)
    
    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    return _to_generated_cv_response(cv_entity)


@router.get("/{cv_id}/versions", response_model=List[GeneratedCVVersionResponse])
async def list_generated_cv_versions(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)
    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    versions = await cv_repo.list_versions(user_id, cv_entity.conversation_id)
    return [_to_generated_cv_version_response(item) for item in versions]


async def _download_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Download generated CV as markdown or docx."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)

    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    stored_format, stored_content = _get_generated_content_payload(cv_entity)
    content_data = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    export_format = format or stored_format

    if export_format not in {"markdown", "docx"}:
        raise HTTPException(status_code=400, detail="Định dạng export không hợp lệ")

    export_content = content_data.get("markdown") or stored_content

    if not str(export_content).strip():
        raise HTTPException(status_code=400, detail="CV không có nội dung để export")

    if export_format == "docx":
        docx_content = _markdown_to_docx_bytes(str(export_content))
        filename = _build_export_filename(cv_entity, "docx")
        return StreamingResponse(
            io.BytesIO(docx_content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if export_format == "markdown":
        filename = _build_export_filename(cv_entity, "md")
        return Response(
            content=str(export_content),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )


@router.get("/{cv_id}/download")
async def download_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await _download_generated_cv(cv_id, format, user_id, session)


@router.get("/{cv_id}/export")
async def export_generated_cv(
    cv_id: UUID,
    format: Literal["markdown", "docx"] | None = Query(default=None),
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await _download_generated_cv(cv_id, format, user_id, session)


@router.post("/{cv_id}/versions", response_model=GeneratedCVResponse)
async def create_generated_cv_version(
    cv_id: UUID,
    req: GeneratedCVUpdateRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Create a new immutable version after user edits in preview."""
    cv_repo = GeneratedCVRepository(session)
    cv_entity = await cv_repo.get_by_id(cv_id)

    if not cv_entity or cv_entity.user_id != user_id:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    existing_payload = cv_entity.generated_content if isinstance(cv_entity.generated_content, dict) else {}
    next_payload = {
        key: value
        for key, value in existing_payload.items()
        if key not in {"html", "import_preview_format"}
    }
    new_entity = await cv_repo.create_versioned(
        user_id=user_id,
        conversation_id=cv_entity.conversation_id,
        parent_version_id=cv_entity.id,
        target_jd_text=cv_entity.target_jd_text,
        base_profile_data=cv_entity.base_profile_data,
        generated_content={
            **next_payload,
            "format": req.output_format,
            "content": req.content,
            "markdown": req.content,
        },
        status=cv_entity.status,
    )

    await session.commit()
    return _to_generated_cv_response(new_entity)


@router.patch("/{cv_id}", response_model=GeneratedCVResponse)
async def update_generated_cv(
    cv_id: UUID,
    req: GeneratedCVUpdateRequest,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    return await create_generated_cv_version(cv_id, req, user_id, session)

@router.delete("/{cv_id}", status_code=204)
async def delete_generated_cv(
    cv_id: UUID,
    user_id: UUID = Depends(get_current_user_id),
    session: AsyncSession = Depends(get_db_session),
):
    """Soft delete generated CV."""
    cv_repo = GeneratedCVRepository(session)
    success = await cv_repo.soft_delete(cv_id, user_id)
    
    if not success:
        raise HTTPException(status_code=404, detail="Không tìm thấy CV mẫu này")

    await session.commit()
